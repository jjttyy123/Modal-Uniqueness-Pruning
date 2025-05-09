# utils/model_eval.py
import torch
import torch.nn as nn
import yaml
import numpy as np
from pathlib import Path
from tqdm import tqdm
from utils.datasets import create_dataloader_rgb_ir
from utils.general import check_dataset, check_img_size, box_iou, non_max_suppression, scale_coords, xyxy2xywh, xywh2xyxy, colorstr
from utils.metrics import ap_per_class, ConfusionMatrix
from utils.plots import output_to_target
from utils.torch_utils import time_synchronized
import json
from docx import Document
from utils.loss import ComputeLoss

def multimodal_eval(model, data_cfg, device, split='val', zero_mode=None,
                    batch_size=32, imgsz=640, conf_thres=0.001, iou_thres=0.6,
                    half_precision=True, verbose=False, opt=None):
    """
    模态缺失模拟评估函数。

    参数：
        model: 要评估的模型（必须支持 RGB + IR 输入）
        data_cfg: 数据配置文件路径或字典，必须包含 'train_rgb', 'train_ir', 'val_rgb', 'val_ir', 'nc', 'names'
        device: 模型推理设备，如 'cuda:0'
        split: 评估数据集划分，选项为 'train', 'val', 'all'
        zero_mode: 模态缺失控制（None 表示无缺失；'rgb' 表示 RGB 缺失；'ir' 表示 IR 缺失）
        batch_size: 批处理大小
        imgsz: 输入图像尺寸
        conf_thres: 置信度阈值
        iou_thres: NMS 阈值
        half_precision: 是否使用半精度推理
        verbose: 是否输出详细信息
        opt: 可选配置对象，用于控制 dataloader 和保存路径

    返回：
        metrics: dict，包括 P, R, mAP@0.5, mAP@0.75, mAP@0.5:0.95
        pred_all: list[Tensor]，所有图像的预测框
        all_train_out: list[Any]，模型 forward 时返回的 train_out 信息
        all_targets: list[Tensor]，原始 ground truth 标签（未缩放）
    """
    assert zero_mode in [None, 'rgb', 'ir'], "zero_mode must be one of None, 'rgb', or 'ir'"
    assert split in ['train', 'val', 'all'], "split must be 'train', 'val', or 'all'"

    # 构建默认配置
    if opt is None:
        class DummyOpt:
            single_cls = False
            rect = True
            cache = False
            stride = 32
            pad = 0.5
            augment = False
            image_weights = False
            workers = 8
            save_txt = False
            save_conf = True
            save_hybrid = False
            project = 'runs/test'
            name = 'exp_eval'
            exist_ok = True
        opt = DummyOpt()

    # 加载数据配置文件
    if isinstance(data_cfg, str):
        with open(data_cfg) as f:
            data_cfg = yaml.safe_load(f)
    check_dataset(data_cfg)
    nc = int(data_cfg['nc'])
    names = data_cfg['names'] if isinstance(data_cfg['names'], dict) else {i: n for i, n in enumerate(data_cfg['names'])}

    # 设置 RGB 和 IR 路径
    if split == 'train':
        rgb_paths, ir_paths = data_cfg['train_rgb'], data_cfg['train_ir']
    elif split == 'val':
        rgb_paths, ir_paths = data_cfg['val_rgb'], data_cfg['val_ir']
    else:
        rgb_paths = data_cfg['train_rgb'] + data_cfg['val_rgb']
        ir_paths = data_cfg['train_ir'] + data_cfg['val_ir']

    # 设置保存目录
    save_dir = Path(opt.project) / opt.name
    (save_dir / 'labels' if opt.save_txt else save_dir).mkdir(parents=True, exist_ok=True)

    # 创建 dataloader
    gs = 32  # grid size
    dataloader = create_dataloader_rgb_ir(rgb_paths, ir_paths, imgsz, batch_size, gs, opt,
                                           pad=0.5, workers=opt.workers, rect=True,
                                           prefix=colorstr(f'[{split.upper()}]: '))[0]

    # 设置模型评估模式及精度
    model.eval()
    half = device.type != 'cpu' and half_precision
    if half:
        model.half()

    # 初始化评估统计量
    iouv = torch.linspace(0.5, 0.95, 10).to(device)
    niou = iouv.numel()
    confusion_matrix = ConfusionMatrix(nc=nc)

    seen, t0, t1 = 0, 0., 0.
    stats, jdict, pred_all, all_targets, all_train_out = [], [], [], [], []
    loss = torch.zeros(3, device=device)
    compute_loss = ComputeLoss(model)

    # 遍历所有的batch数据
    for batch_i, (img, targets, paths, shapes) in enumerate(tqdm(dataloader, desc='Evaluating')):
        # 将图像数据移到设备（如GPU）并进行非阻塞传输
        img = img.to(device, non_blocking=True)
        # 如果需要，使用半精度浮点数（FP16），否则使用单精度浮点数（FP32）
        img = img.half() if half else img.float()
        # 将图像像素值缩放到[0, 1]范围
        img /= 255.0
        # 将目标标签移到设备上
        targets = targets.to(device)

        # 获取当前批次图像的数量、通道数、高度和宽度
        nb, _, h, w = img.shape

        # 将图像数据分为RGB和IR（红外）两个部分
        # img 是一个形状为 (nb, c, h, w) 的张量，这里是[64, 6, 544, 672]
        # 第一个":"表示选择批次中的所有图像（batch size）。
        # ":3"表示选择每张图像的前 3 个通道，即红色、绿色和蓝色通道。
        img_rgb, img_ir = img[:, :3], img[:, 3:]
        # print(img.size())

        # 根据指定的零模式（zero_mode）模拟模态的缺失
        if zero_mode == 'rgb':
            img_rgb.zero_()  # 将RGB图像数据置零
        elif zero_mode == 'ir':
            img_ir.zero_()  # 将IR图像数据置零

        # 不需要计算梯度（推理阶段），避免浪费内存和计算资源
        with torch.no_grad():
            # 记录模型推理时间
            t = time_synchronized()
            # 执行模型的前向推理，输出预测结果和训练阶段的中间输出
            out, train_out = model(img_rgb, img_ir, augment=opt.augment)
            # 计算前向推理时间并累加
            t0 += time_synchronized() - t

            # Compute loss
            if compute_loss:
                loss += compute_loss([x.float() for x in train_out], targets)[1][:3]  # box, obj, cls

            # 保存训练阶段的输出（中间结果）
            all_train_out.append(train_out)
            # 保存原始标签（没有进行缩放）
            all_targets.append(targets.clone())

            # 对标签框进行缩放，使其与图像的尺寸匹配
            targets[:, 2:] *= torch.Tensor([w, h, w, h]).to(device)
            # 如果需要，保存混合标签（Hybrid Label）
            lb = [targets[targets[:, 0] == i, 1:] for i in range(nb)] if opt.save_hybrid else []

            # 记录非最大抑制的时间
            t = time_synchronized()
            # 对模型输出进行非最大抑制（NMS），过滤掉低置信度的框，并进行IoU阈值处理
            out = non_max_suppression(out, conf_thres, iou_thres, labels=lb,
                                    multi_label=True, agnostic=opt.single_cls)
            # 计算NMS操作的时间并累加
            t1 += time_synchronized() - t

        # 对每一张图片逐个计算评估指标
        for si, pred in enumerate(out):
            # 获取当前图片的真实标签
            labels = targets[targets[:, 0] == si, 1:]
            nl = len(labels)  # 标签的数量
            tcls = labels[:, 0].tolist() if nl else []  # 标签的类别列表
            seen += 1  # 总共处理的图片数量

            # 如果没有预测结果，继续处理下一个batch
            if len(pred) == 0:
                # 如果有标签，则保存空的评估数据（未检测到目标）
                if nl:
                    stats.append((torch.zeros(0, niou, dtype=torch.bool), torch.Tensor(), torch.Tensor(), tcls))
                continue

            # 对预测框进行克隆，以免修改原始数据
            predn = pred.clone()
            # 将预测框的坐标从图像的原始尺寸缩放到实际的尺寸
            scale_coords(img[si].shape[1:], predn[:, :4], shapes[si][0], shapes[si][1])
            pred_all.append(predn)  # 保存所有预测框

            # 如果是单类别任务，强制将预测类别设置为0
            if opt.single_cls:
                pred[:, 5] = 0

            # 创建一个布尔矩阵来记录每个预测框是否正确（通过IoU进行评估）
            correct = torch.zeros(pred.shape[0], niou, dtype=torch.bool, device=device)
            if nl:
                # 将标签的框从xywh格式转换为xyxy格式
                tbox = xywh2xyxy(labels[:, 1:5])
                # 将标签框的坐标从图像的原始尺寸缩放到实际的尺寸
                scale_coords(img[si].shape[1:], tbox, shapes[si][0], shapes[si][1])

                # 对每一个类别，计算IoU并匹配预测框和标签框
                for cls in torch.unique(labels[:, 0]):
                    # 获取当前类别的标签框
                    ti = (cls == labels[:, 0]).nonzero(as_tuple=False).view(-1)
                    # 获取当前类别的预测框
                    pi = (cls == pred[:, 5]).nonzero(as_tuple=False).view(-1)
                    if pi.shape[0]:
                        # 计算IoU，找出最匹配的预测框和标签框
                        ious, i = box_iou(predn[pi, :4], tbox[ti]).max(1)
                        detected_set = set()
                        # 根据IoU值过滤掉过低的匹配
                        for j in (ious > iouv[0]).nonzero(as_tuple=False):
                            d = ti[i[j]]
                            if d.item() not in detected_set:
                                detected_set.add(d.item())
                                correct[pi[j]] = ious[j] > iouv
                                if len(detected_set) == nl:
                                    break
            # 将预测框、置信度、类别等信息保存下来
            stats.append((correct.cpu(), pred[:, 4].cpu(), pred[:, 5].cpu(), tcls))
        #     print("tp:")
        #     print(correct)
        #     print("conf:")
        #     print(pred[:, 4])
        #     print("pred_cls")
        #     print(pred[:, 5])
        #     print("target_cls:")
        #     print(tcls)
        # exit(0)



    # 汇总评估指标
    stats = [np.concatenate(x, 0) for x in zip(*stats)]
    if len(stats) and stats[0].any():
        p, r, ap, f1, ap_class = ap_per_class(*stats, plot=False, save_dir=save_dir, names=names)
        ap50, ap75, ap = ap[:, 0], ap[:, 5], ap.mean(1)
        mp, mr, map50, map75, map = p.mean(), r.mean(), ap50.mean(), ap75.mean(), ap.mean()
        nt = np.bincount(stats[3].astype(np.int64), minlength=nc)  # number of targets per class
    else:
        nt = torch.zeros(1)
        mp = mr = map50 = map75 = map = 0.0

    # 打印头部信息
    header = ('%20s' + '%12s' * 7) % ('Class', 'Images', 'Labels', 'P', 'R', 'mAP@.5', 'mAP@.75', 'mAP@.5:.95')
    print(header)

    # 打印数据
    pf = '%20s' + '%12i' * 2 + '%12.3g' * 5  # 数据的格式化方式
    print(pf % ('all', seen, nt.sum(), mp, mr, map50, map75, map))

    # Print results per class
    if nc > 1 and len(stats):
        for i, c in enumerate(ap_class):
            print(pf % (names[c], seen, nt[c], p[i], r[i], ap50[i], ap75[i], ap[i]))
    
    # 计算真实标签p_y以及预测标签p_y_x
    p_y_x,p_y = compute_pyx_py(*stats, plot=False, save_dir=save_dir, names=names)

    return ap_class,p_y_x,p_y,len(dataloader)

def compute_pyx_py(tp, conf, pred_cls, target_cls, plot=False, save_dir='.', names=()):

    """
    计算每个真实框对应的预测概率（p_y_x）和真实标签概率（p_y）
    只使用 mAP50（即 tp[:, 0]） 为 True 的预测框进行处理
    """

    # 1、预处理
    # 1.1 排序预测框的置信度，按置信度降序排列
    i = np.argsort(-conf)
    tp, conf, pred_cls = tp[i], conf[i], pred_cls[i]

    # 1.2 Find unique classes
    unique_classes = np.unique(target_cls)
    nc = unique_classes.shape[0]  # 类别数量

    # 1.3 只使用 IoU=0.5 匹配成功的预测框进行处理，提高效率
    valid_mask = tp[:, 0] == True
    tp, conf, pred_cls = tp[valid_mask], conf[valid_mask], pred_cls[valid_mask]

    # 2、初始化存储预测类别概率（p_y_x）和真实标签概率（p_y）的数组
    num_gt = len(target_cls)  # 真实框数量
    p_y_x = np.zeros((num_gt, nc), dtype=float) # 每个预测框对应的预测单热编码
    p_y = np.zeros((num_gt, nc), dtype=float)   # 每个真实框对应的预测单热编码
    matched = [False] * len(pred_cls)  # 匹配状态标记（每个预测框只能匹配一次）
    cnt_not_process = 0

    for c in unique_classes:
        target_idx = np.where(target_cls == c)[0]  # 真实框索引 
        pred_idx = np.where(pred_cls == c)[0]      # 预测框索引
        print(f"c:{c} len(target_idx):{len(target_idx)}")
        print(f"c:{c} len(pred_idx):{len(pred_idx)}")

        for idx, true_idx in enumerate(target_idx):
            true_cls = int(target_cls[true_idx])
            p_y[true_idx][true_cls] = 1    # 真实类别置1，其它类别为0

            # 寻找匹配的预测框
            matched_pred_idx = None  # 记录匹配的预测框索引
            for pred_id in pred_idx:
                if not matched[pred_id]:
                    # 正类赋值为置信度
                    p_y_x[true_idx][true_cls] = conf[pred_id]
                    matched_pred_idx = pred_id
                    matched[pred_id] = True
                    break  # 每个真实框只需匹配一个预测框

            # 如果没有找到匹配的预测框，p_y_x 保持为 0，也就是不用管
            if matched_pred_idx is None:
                cnt_not_process += 1

    print(f"[compute_pyx_py] 未匹配的真实框数: {cnt_not_process}/{num_gt} ({cnt_not_process / (num_gt + 1e-16):.2%})")
    
    # 把numpy.ndarray转为tensor
    p_y_x = torch.tensor(p_y_x)
    p_y = torch.tensor(p_y)

    # 对于单个类别，需要补全一下，不然互信息会出现负值
    if(nc == 1):
        p_y_x = torch.cat([p_y_x, torch.zeros_like(p_y_x)], dim=1)  # 补全 p_y_x1 的第二列
        p_y = torch.cat([p_y, torch.zeros_like(p_y)])  # 补全 p_y 的第二列

    return p_y_x, p_y



def eval_all_combination(model,model_only_RGB,model_only_IR,
                        data_cfg,device,
                        json_path='./evaluation_results.json'):
    # 循环评估
    splits = ['train', 'val']
    zero_modes = [None, 'rgb', 'ir']
    models = {
        'normal': model,
        'IR_zero': model_only_RGB,
        'RGB_zero': model_only_IR
    }

    # 存储结果的字典
    results = {}

    # 循环评估所有18种情况
    for split in splits:
        results[split] = {}
        for model_key, current_model in models.items():
            results[split][model_key] = {}
            for zero_mode in zero_modes:
                # 当前评估配置
                print(f"Evaluating: Split={split}, Model={model_key}, Zero_Mode={zero_mode}")
                
                # 调用评估函数
                ap_class,p_y_x,p_y = multimodal_eval(
                    model=current_model,
                    data_cfg=data_cfg,
                    device=device,
                    split=split,
                    zero_mode=zero_mode,
                    batch_size=64
                )
                # 存储评估结果
                zero_mode_key = zero_mode if zero_mode else 'none'
                results[split][model_key][zero_mode_key] = ap_class

    # 将结果保存到JSON文件
    with open(json_path, 'w') as f:
        json.dump(results, f, indent=4)

    # 打印评估完成信息
    print(f"所有评估完成，结果已保存到{json_path}")

def json_to_word(json_path='./evaluation_results.json',word_path='./evaluation_results.docx'):
    # 读取JSON数据
    with open(json_path, 'r') as f:
        data = json.load(f)

    # 创建Word文档
    doc = Document()
    doc.add_heading('Evaluation Results', 0)

    # 循环数据创建表格
    for split, models in data.items():
        doc.add_heading(f'Split: {split}', level=1)

        for model_name, zero_modes in models.items():
            doc.add_heading(f'Model: {model_name}', level=2)

            # 创建表格，表头
            table = doc.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Zero Mode'
            hdr_cells[1].text = 'Precision'
            hdr_cells[2].text = 'Recall'
            hdr_cells[3].text = 'mAP@0.5'
            hdr_cells[4].text = 'mAP@0.75'
            hdr_cells[5].text = 'mAP@0.5:0.95'

            # 填充表格数据
            for zero_mode, metrics in zero_modes.items():
                row_cells = table.add_row().cells
                row_cells[0].text = zero_mode
                row_cells[1].text = f"{metrics['P']:.4f}"
                row_cells[2].text = f"{metrics['R']:.4f}"
                row_cells[3].text = f"{metrics['mAP@0.5']:.4f}"
                row_cells[4].text = f"{metrics['mAP@0.75']:.4f}"
                row_cells[5].text = f"{metrics['mAP@0.5:0.95']:.4f}"

            doc.add_paragraph()  # 添加空行以区分不同的模型

    # 保存Word文档
    doc.save(word_path)
    print(f"JSON 数据已成功转换为 {word_path} 文件")
